from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.services.report.report_service import ReportService
from app.services.report.error_notebook_service import ErrorNotebookService
from app.services.report.practice_generator_service import PracticeGeneratorService
import io
import urllib.parse
from collections import defaultdict
from typing import List, Optional
from pydantic import BaseModel

router = APIRouter(prefix="/report", tags=["报告导出"])


def get_service(db: Session) -> ReportService:
    return ReportService(db)


@router.get("/word/{exam_id}")
def export_word(exam_id: str, db: Session = Depends(get_db)):
    try:
        service = get_service(db)
        buf = service.generate_word(exam_id)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=analysis_report.docx"},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/pdf/{exam_id}")
def export_pdf(exam_id: str, db: Session = Depends(get_db)):
    try:
        service = get_service(db)
        buf = service.generate_pdf(exam_id)
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=analysis_report.pdf"},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/ppt/{exam_id}")
def export_ppt(exam_id: str, db: Session = Depends(get_db)):
    try:
        service = get_service(db)
        buf = service.generate_ppt(exam_id)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            headers={"Content-Disposition": f"attachment; filename=analysis_report.pptx"},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/word/class/{class_id}")
def export_class_word(class_id: str, db: Session = Depends(get_db)):
    try:
        service = get_service(db)
        buf = service.generate_class_word(class_id)
        if not buf:
            raise HTTPException(status_code=404, detail="班级数据不存在")
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=class_analysis_report.docx"},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/pdf/class/{class_id}")
def export_class_pdf(class_id: str, db: Session = Depends(get_db)):
    try:
        service = get_service(db)
        buf = service.generate_class_pdf(class_id)
        if not buf:
            raise HTTPException(status_code=404, detail="班级数据不存在")
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=class_analysis_report.pdf"},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/word/student/{student_id}")
def export_student_word(student_id: str, db: Session = Depends(get_db)):
    try:
        service = get_service(db)
        buf = service.generate_student_word(student_id)
        if not buf:
            raise HTTPException(status_code=404, detail="学生数据不存在")
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=student_analysis_report.docx"},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/pdf/student/{student_id}")
def export_student_pdf(student_id: str, db: Session = Depends(get_db)):
    try:
        service = get_service(db)
        buf = service.generate_student_pdf(student_id)
        if not buf:
            raise HTTPException(status_code=404, detail="学生数据不存在")
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=student_analysis_report.pdf"},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/excel/{exam_id}")
def export_excel(exam_id: str, db: Session = Depends(get_db)):
    from app.services.analytics.analysis_service import AnalysisService
    import io, openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from fastapi.responses import StreamingResponse
    import urllib.parse
    
    service = AnalysisService(db)
    data = service.get_exam_analysis(exam_id)
    if not data:
        raise HTTPException(status_code=404, detail="考试不存在")
    
    wb = openpyxl.Workbook()
    
    # Sheet 1: 年级统计
    ws = wb.active
    ws.title = "年级统计"
    hf = Font(bold=True, color="FFFFFF", size=11)
    hfill = PatternFill(start_color="409EFF", end_color="409EFF", fill_type="solid")
    ha = Alignment(horizontal="center", vertical="center")
    tb = Border(left=Side(style="thin"), right=Side(style="thin"), top=Side(style="thin"), bottom=Side(style="thin"))
    
    headers = ["科目", "平均分", "得分率", "最高分", "最低分", "及格率", "优秀率", "标准差"]
    ws.append(["考试名称", data.get("exam_name", "")])
    ws.append(["考试日期", data.get("exam_date", "")])
    ws.append(["参考人数", data.get("total_students", 0)])
    ws.append([])
    ws.append(headers)
    for ci in range(1, len(headers)+1):
        c = ws.cell(row=5, column=ci)
        c.font, c.fill, c.alignment, c.border = hf, hfill, ha, tb
    
    for gs in data.get("grade_stats", []):
        ws.append([gs.get("subject_name", ""), gs.get("avg_score", ""), f"{gs.get('avg_score_rate', '')}%", gs.get("max_score", ""), gs.get("min_score", ""), f"{gs.get('pass_rate', '')}%", f"{gs.get('excellent_rate', '')}%", gs.get("std_dev", "")])
    
    ws.column_dimensions["A"].width = 14
    for c in "BCDEFGH":
        ws.column_dimensions[c].width = 12
    
    # Sheet 2: 班级统计
    ws2 = wb.create_sheet("班级统计")
    class_stats = data.get("class_stats", [])
    if class_stats:
        all_subj = list(set(s.get("subject_name", "") for cs in class_stats for s in cs.get("stats", [])))
        h2 = ["班级", "人数"] + all_subj
        ws2.append(h2)
        for ci in range(1, len(h2)+1):
            c = ws2.cell(row=1, column=ci)
            c.font, c.fill, c.alignment, c.border = hf, hfill, ha, tb
        for cs in class_stats:
            row_data = [cs.get("class_name", ""), cs.get("student_count", 0)]
            smap = {s.get("subject_name", ""): s.get("avg_score", "") for s in cs.get("stats", [])}
            for subj in all_subj:
                row_data.append(smap.get(subj, ""))
            ws2.append(row_data)
    
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={urllib.parse.quote(data.get('exam_name', 'report') + '_分析报表.xlsx')}"},
    )


class PracticeGenerateRequest(BaseModel):
    """生成针对性练习请求"""
    question_count: Optional[int] = 10
    include_types: Optional[List[str]] = None


def build_error_notebook_docx(data: dict) -> io.BytesIO:
    """生成错题集 Word 文档"""
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    title = doc.add_heading(f"{data.get('student_name', '')} 错题集", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"错题总数: {data.get('total_errors', 0)}")

    sections = defaultdict(list)
    for err in data.get("errors", []):
        sections[err.get("subject_name") or "未分类"].append(err)

    for subject_name, errors in sections.items():
        doc.add_heading(subject_name, level=1)
        for idx, err in enumerate(errors, 1):
            doc.add_paragraph(
                f"{idx}. 第 {err.get('question_no', '-')} 题"
                f"[{err.get('question_type', '')}] "
                f"知识点: {err.get('knowledge_point_name', '') or '未标注'}"
            )
            if err.get("question_content"):
                doc.add_paragraph(f"题目: {err['question_content']}")
            doc.add_paragraph(
                f"得分: {err.get('score_earned', 0)} / {err.get('full_score', 0)}"
                f"    失分率: {err.get('loss_rate', 0)}%"
            )
            doc.add_paragraph("")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


@router.get("/error-notebook/{student_id}")
def get_error_notebook(student_id: str, exam_id: str = None, db: Session = Depends(get_db)):
    """获取学生错题集"""
    try:
        service = ErrorNotebookService(db)
        data = service.generate_error_notebook(student_id, exam_id)
        sections = defaultdict(list)
        for err in data.get("errors", []):
            subj = err.get("subject_name") or "未分类"
            sections[subj].append({
                "question_no": err.get("question_no"),
                "question_type": err.get("question_type", ""),
                "knowledge_point": err.get("knowledge_point_name", "") or "",
                "content": err.get("question_content", "") or "",
                "your_score": err.get("score_earned"),
                "full_score": err.get("full_score"),
                "loss_rate": err.get("loss_rate", 0),
                "difficulty": err.get("difficulty"),
                "exam_name": err.get("exam_name", "") or "",
                "analysis": f"本题失分率 {err.get('loss_rate', 0)}%，建议加强该知识点练习",
            })
        return {
            "student_name": data.get("student_name", ""),
            "total_errors": data.get("total_errors", len(data.get("errors", []))),
            "sections": [{"subject_name": subj, "items": items} for subj, items in sections.items()],
            "knowledge_summary": data.get("knowledge_summary", []),
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/error-notebook/{student_id}/export")
def export_error_notebook(student_id: str, exam_id: str = None, db: Session = Depends(get_db)):
    """导出学生错题集 Word"""
    try:
        service = ErrorNotebookService(db)
        data = service.generate_error_notebook(student_id, exam_id)
        if not data.get("errors"):
            raise HTTPException(status_code=404, detail="暂无错题数据")
        buf = build_error_notebook_docx(data)
        filename = urllib.parse.quote((data.get("student_name") or "student") + "_错题集.docx")
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/practice/{student_id}")
def generate_student_practice(student_id: str, payload: PracticeGenerateRequest = None, db: Session = Depends(get_db)):
    """生成针对性练习"""
    try:
        question_count = payload.question_count if payload else 10
        include_types = payload.include_types if payload else None
        from app.services.analytics.knowledge_analysis_service import KnowledgeAnalysisService
        kp_service = KnowledgeAnalysisService(db)
        analysis = kp_service.get_student_knowledge_analysis(student_id)
        weaknesses = analysis.get("weaknesses", []) if analysis else []
        practice_service = PracticeGeneratorService(db)
        data = practice_service.generate_practice(student_id, weaknesses, question_count, include_types)
        questions = []
        for sheet in data.get("practice_sheets", []):
            for q in sheet.get("questions", []):
                questions.append({
                    "type": q.get("question_type", ""),
                    "difficulty": q.get("difficulty", ""),
                    "content": q.get("content", ""),
                    "knowledge_point": q.get("knowledge_point", ""),
                    "hint": q.get("hint", ""),
                    "estimated_time": q.get("estimated_time", ""),
                })
        return {
            "student_id": student_id,
            "title": data.get("message") or "个性化提升练习",
            "questions": questions,
            "total_questions": data.get("total_questions", len(questions)),
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

