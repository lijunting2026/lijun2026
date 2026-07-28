"""Word document generation for exam reports."""
import io
from typing import Dict, Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _add_chart_to_docx(doc, chart_buf, width_inches=5.5):
    """Add a chart image to a Word document"""
    if chart_buf and chart_buf.getbuffer().nbytes > 1000:
        doc.add_picture(chart_buf, width=Inches(width_inches))
        doc.add_paragraph("")


def generate_class_word(class_id: str, db, analysis_service) -> io.BytesIO:
    """Generate Word report for a class."""
    from app.services.analytics.student_tracking import ClassAnalysisService
    service = ClassAnalysisService(db)
    data = service.get_class_overview(class_id)
    if not data:
        return None

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    title = doc.add_heading(f"{data.get('class_name', '')} 班级分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"年级: {data.get('grade_name', '-')}    学生人数: {data.get('student_count', 0)}    考试次数: {data.get('exam_count', 0)}")
    doc.add_paragraph("")

    doc.add_heading("各科统计", level=1)
    subject_stats = data.get("subject_stats", [])
    if subject_stats:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["科目", "平均分", "最高分", "最低分", "样本数"]):
            hdr[i].text = h
        for ss in subject_stats:
            row = table.add_row().cells
            row[0].text = ss.get("subject_name", "")
            row[1].text = str(ss.get("avg_score", ""))
            row[2].text = str(ss.get("max_score", ""))
            row[3].text = str(ss.get("min_score", ""))
            row[4].text = str(ss.get("count", ""))

    doc.add_paragraph("")
    doc.add_heading("历次考试", level=1)
    for es in data.get("exam_summary", []):
        doc.add_paragraph(f"{es.get('exam_name', '')} - 平均得分率: {es.get('avg_rate', '')}%")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_student_word(student_id: str, db, analysis_service) -> io.BytesIO:
    """Generate Word report for a student."""
    from app.services.analytics.student_tracking import StudentTrackingService
    service = StudentTrackingService(db)
    data = service.get_student_scores(student_id)
    if not data:
        return None

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    title = doc.add_heading(f"{data.get('student_name', '')} 学情跟踪报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("成绩概览", level=1)
    doc.add_paragraph(f"考试次数: {data.get('exam_count', 0)}")
    doc.add_paragraph(f"总体趋势: {data.get('overall_trend', '')}")
    if data.get("strengths"):
        doc.add_paragraph("优势科目: " + ", ".join([s['subject_name'] for s in data['strengths']]))
    if data.get("weaknesses"):
        doc.add_paragraph("薄弱科目: " + ", ".join([s['subject_name'] for s in data['weaknesses']]))

    doc.add_heading("历次考试详情", level=1)
    for exam in data.get("exams", []):
        doc.add_heading(f"{exam.get('exam_name', '')} ({exam.get('exam_date', '')})", level=2)
        doc.add_paragraph(f"平均得分率: {exam.get('avg_rate', '')}%")
        subjects_table = doc.add_table(rows=1, cols=3)
        subjects_table.style = 'Light Grid Accent 1'
        hdr = subjects_table.rows[0].cells
        for i, h in enumerate(["科目", "分数", "得分率"]):
            hdr[i].text = h
        for s in exam.get("subjects", []):
            row = subjects_table.add_row().cells
            row[0].text = s.get("subject_name", "")
            row[1].text = str(s.get("score", ""))
            row[2].text = f"{s.get('rate', '')}%"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_exam_word(exam_id: str, db, analysis_service) -> io.BytesIO:
    """Generate Word report for an exam."""
    data = analysis_service.get_exam_analysis(exam_id)
    if not data:
        return None

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    title = doc.add_heading(f"{data.get('exam_name', '')} 质量分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"考试日期: {data.get('exam_date', '-')}    参考人数: {data.get('total_students', 0)}")
    doc.add_paragraph("")

    doc.add_heading("年级统计", level=1)
    grade_stats = data.get("grade_stats", [])
    if grade_stats:
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        for i, h in enumerate(["科目", "平均分", "得分率%", "最高分", "最低分", "及格率%", "优秀率%"]):
            hdr[i].text = h
        for gs in grade_stats:
            row = table.add_row().cells
            row[0].text = gs.get("subject_name", "")
            row[1].text = str(gs.get("avg_score", ""))
            row[2].text = str(gs.get("avg_score_rate", ""))
            row[3].text = str(gs.get("max_score", ""))
            row[4].text = str(gs.get("min_score", ""))
            row[5].text = str(gs.get("pass_rate", ""))
            row[6].text = str(gs.get("excellent_rate", ""))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf